import docx
import re
import os
from docx.shared import Inches
from typing import List

def createDoc(quesfileName: str,allImages: List[str]):
    try:
        file = open(quesfileName,'r')
        fileList = list(file)
        questions = list(filter(lambda x: True if(x!='\n') else False,fileList))
        doc = docx.Document()
        imageIdx = 0
        print(allImages)
        for x in range(len(questions)):
            quesParaObj = doc.add_paragraph('Question: '+str(x+1)+":")
            quesParaObj.add_run('\n')
            quesParaObj.add_run(questions[x])
            while(imageIdx<len(allImages) and re.split("[_.]",allImages[imageIdx])[1]==str(x+1)):
                # print(allImages[imageIdx],imageIdx)
                doc.add_picture('image/'+allImages[imageIdx],width=Inches(5))
                imageIdx += 1

        doc.save('assignment.docx')
        return True
    except:
        return False
      asdsa

def getImagesName() -> List[str]:
    # Get Image folder path
    imagesPath = os.getcwd()+"/image"
    
    # Get all Images names
    allImages = os.listdir(imagesPath)

    # Sort Images according to question number
    allImages.sort(key= lambda x: (int(re.split("[_.]",x)[1]), int(re.split("[_.]",x)[2])))

    return allImages

def main():
    if createDoc('questionFile',getImagesName()):
        print("Doc Created")
    else:
        print("Something Went Wrong")
    
main()